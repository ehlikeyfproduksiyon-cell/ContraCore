# -*- coding: utf-8 -*-
"""
Word dosyası okuma ve Adım-5 Excel üretimi — saf Python, yan etkisiz.
Selenium veya Playwright importu yoktur; bağımsız test edilebilir.
"""
import os
import re
from pathlib import Path


# ══════════════════════════════════════════════════════════
#  WORD OKUYUCU
# ══════════════════════════════════════════════════════════

def doc_oku(filepath: str) -> dict | None:
    """
    .doc veya .docx dosyasını okur.
    Döner: {'html': str|None, 'paragraflar': list, 'tablolar': list}
    Hata durumunda None.
    """
    filepath = str(filepath)
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".docx":
        try:
            import docx as docx_lib
            doc = docx_lib.Document(filepath)
            paragraflar = [p.text for p in doc.paragraphs]
            tablolar = []
            for tbl in doc.tables:
                for row in tbl.rows:
                    tablolar.append([c.text.strip() for c in row.cells])
            return {"paragraflar": paragraflar, "tablolar": tablolar, "html": None}
        except Exception as e:
            print(f"[UYARI] docx okuma hatası: {e}")
            return None

    # .doc — HTML tabanlı mı kontrol et
    try:
        with open(filepath, "rb") as f:
            header = f.read(10)
    except OSError as e:
        print(f"[UYARI] Dosya açılamadı: {e}")
        return None

    # UTF-16 LE BOM kontrolü (FF FE) — Word bazen UTF-16 kaydeder
    if header[:2] == b'\xff\xfe':
        try:
            with open(filepath, "r", encoding="utf-16-le", errors="replace") as f:
                html = f.read()
            return {"html": html, "paragraflar": [], "tablolar": []}
        except Exception:
            pass

    # UTF-16 BE BOM kontrolü (FE FF)
    if header[:2] == b'\xfe\xff':
        try:
            with open(filepath, "r", encoding="utf-16-be", errors="replace") as f:
                html = f.read()
            return {"html": html, "paragraflar": [], "tablolar": []}
        except Exception:
            pass

    if b"<!DOCTYPE" in header or b"<html" in header or b"<!D" in header:
        for enc in ["utf-8", "cp1254", "iso-8859-9", "latin-1"]:
            try:
                with open(filepath, "r", encoding=enc, errors="replace") as f:
                    html = f.read()
                return {"html": html, "paragraflar": [], "tablolar": []}
            except Exception:
                continue

    print(f"[UYARI] Bilinmeyen .doc formatı: {filepath}")
    return None


def html_tablolari_cek(html: str) -> list[list[list[str]]]:
    """HTML içinden tüm tabloları [[satir, ...], ...] olarak döner."""
    _ENTITY = {
        "&nbsp;": " ", "&amp;": "&", "&uuml;": "ü", "&Uuml;": "Ü",
        "&ccedil;": "ç", "&Ccedil;": "Ç", "&ouml;": "ö", "&Ouml;": "Ö",
        "&iacute;": "ı", "&Iacute;": "İ", "&szlig;": "ş", "&hellip;": "…",
    }
    tables = []
    for tbl_html in re.findall(r"<table[^>]*>(.*?)</table>", html, re.DOTALL | re.IGNORECASE):
        rows = []
        for row_html in re.findall(r"<tr[^>]*>(.*?)</tr>", tbl_html, re.DOTALL | re.IGNORECASE):
            cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, re.DOTALL | re.IGNORECASE)
            temiz = []
            for c in cells:
                c = re.sub(r"<[^>]+>", "", c)
                for ent, rep in _ENTITY.items():
                    c = c.replace(ent, rep)
                c = c.replace("\r", " ").replace("\n", " ")
                c = re.sub(r"\s+", " ", c).strip()
                temiz.append(c)
            if any(temiz):
                rows.append(temiz)
        if rows:
            tables.append(rows)
    return tables


def fatura_tablosunu_bul(tablolar: list) -> list | None:
    """'Fatura Tarihi' ve 'Fatura No' içeren tabloyu döner."""
    for tablo in tablolar:
        for satir in tablo:
            metin = " ".join(satir).lower()
            if "fatura tarihi" in metin and ("fatura no" in metin or "fatura n" in metin):
                return tablo
    return None


def tutar_parse(metin: str) -> float | str:
    """'83.400,00' → 83400.0. Başarısız → boş string."""
    if not metin:
        return ""
    metin = re.sub(r"[^\d.,]", "", str(metin).strip())
    if not metin:
        return ""
    try:
        return float(metin.replace(".", "").replace(",", "."))
    except Exception:
        return ""


def faturalari_cek(word_yolu: str) -> list[dict]:
    """
    Word dosyasından fatura listesi çıkarır.
    Döner: [{'tarih':..., 'seri':..., 'no':..., 'tutar':..., 'kdv':...}, ...]
    """
    veri = doc_oku(word_yolu)
    if veri is None:
        return []

    if veri["html"]:
        tablolar = html_tablolari_cek(veri["html"])
    else:
        tablolar = [veri["tablolar"]] if veri["tablolar"] else []

    fatura_tbl = fatura_tablosunu_bul(tablolar)
    if not fatura_tbl:
        return []

    # Başlık satırını ve sütun indekslerini bul
    header_idx = None
    kolon: dict[str, int] = {}
    for i, satir in enumerate(fatura_tbl):
        metin = " ".join(satir).lower()
        if "fatura tarihi" in metin:
            header_idx = i
            for j, h in enumerate(satir):
                h_low = h.lower()
                if "tarih" in h_low and j == 0:
                    kolon["tarih"] = j
                elif "no" in h_low and j > 0 and "fatura" in " ".join(satir[:j+1]).lower():
                    kolon.setdefault("no", j)
                elif "matrah" in h_low:
                    kolon["matrah"] = j
                elif "kdv" in h_low and "tutarı" in h_low:
                    kolon["kdv"] = j
            break

    if header_idx is None:
        return []

    kolon.setdefault("tarih",  0)
    kolon.setdefault("no",     1)
    kolon.setdefault("matrah", 3)
    kolon.setdefault("kdv",    4)

    faturalar = []
    for satir in fatura_tbl[header_idx + 1:]:
        metin = " ".join(satir).lower()
        if "toplam" in metin:
            continue
        if len(satir) < 2 or not any(s.strip() for s in satir):
            continue
        no_idx = kolon["no"]
        if no_idx >= len(satir) or not satir[no_idx].strip():
            continue

        fatura_no    = satir[kolon["no"]].strip()    if kolon["no"]    < len(satir) else ""
        fatura_tarih = satir[kolon["tarih"]].strip() if kolon["tarih"] < len(satir) else ""
        matrah       = tutar_parse(satir[kolon["matrah"]]) if kolon["matrah"] < len(satir) else ""
        kdv          = tutar_parse(satir[kolon["kdv"]])    if kolon["kdv"]    < len(satir) else ""

        if fatura_no and fatura_tarih:
            faturalar.append({
                "tarih": fatura_tarih,
                "seri":  "",
                "no":    fatura_no,
                "tutar": matrah,
                "kdv":   kdv,
            })

    return faturalar


def fatura_aylarini_al(faturalar: list[dict]) -> list[str]:
    """
    Fatura listesinden benzersiz, sıralı ay listesi çıkarır.
    Döner: ['03.2025', '07.2025']
    """
    aylar: set[str] = set()
    for f in faturalar:
        tarih = str(f.get("tarih", "")).strip()
        for sep in [".", "/"]:
            parcalar = tarih.split(sep)
            if len(parcalar) == 3:
                try:
                    ay  = int(parcalar[1])
                    yil = int(parcalar[2])
                    if 1 <= ay <= 12 and 2000 <= yil <= 2100:
                        aylar.add(f"{ay:02d}.{yil}")
                except Exception:
                    pass
                break
    return sorted(aylar, key=lambda x: (int(x[3:]), int(x[:2])))


# ══════════════════════════════════════════════════════════
#  ADIM-5 EXCEL ÜRETİCİ
# ══════════════════════════════════════════════════════════

def adim5_excel_uret(faturalar: list[dict], cikti_yolu: str) -> str:
    """GIB Adım-5 formatında Excel dosyası üretir. Oluşturulan yolu döner."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Veriler"

    basliklar = [
        "Faturanın Tarihi",
        "Faturanın Serisi",
        "Faturanın Numarası",
        "Faturanın Tutarı (TL)",
        "K.D.V(TL)",
        "Defter Kayıt Tarihi",
        "Yevmiye Numarası",
        "Ödeme Şekli ve Ödemeye İlişkin Bilgiler",
        "Açıklama",
        "Hatalı Satır Açıklama",
    ]

    ince   = Side(style="thin", color="CCCCCC")
    border = Border(left=ince, right=ince, top=ince, bottom=ince)

    for col, baslik in enumerate(basliklar, 1):
        h = ws.cell(row=1, column=col, value=baslik)
        h.font      = Font(name="Calibri", bold=True, size=10, color="FFFFFF")
        h.fill      = PatternFill("solid", fgColor="1B5E8E")
        h.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        h.border    = border

    ws.row_dimensions[1].height = 30

    for i, f in enumerate(faturalar, 2):
        degerler = [f["tarih"], f["seri"], f["no"], f["tutar"], f["kdv"], "", "", "", "", ""]
        for col, val in enumerate(degerler, 1):
            c = ws.cell(row=i, column=col, value=val)
            c.font      = Font(name="Calibri", size=10)
            c.border    = border
            c.alignment = Alignment(vertical="center")

    genislikler = [14, 10, 25, 18, 14, 14, 14, 30, 20, 20]
    for i, g in enumerate(genislikler, 1):
        ws.column_dimensions[ws.cell(1, i).column_letter].width = g

    wb.save(cikti_yolu)
    return cikti_yolu


# ══════════════════════════════════════════════════════════
#  TASDİK BİLGİSİ ÇIKARTICI
# ══════════════════════════════════════════════════════════

def tasdik_bilgisi_cek(word_yolu: str) -> dict:
    """
    Word dosyasından tasdik ve sözleşme bilgilerini çıkarır.
    Döner: {
        'karsit_firma_adi', 'karsit_vkn', 'mukellef_vkn',
        'kdv_soz_tarih', 'kdv_soz_no',
        'tam_soz_tarih', 'tam_soz_no',
        'tam_tasdik_var': bool,
        'inceleme_donemi',
    }
    """
    # Bilgi İsteme tespiti — sadece "Konu" alanında geçenler
    # NOT: "Hakkında Bilgi İstenen Fatura" gibi standart tutanak başlıkları hariç
    _BILGI_ISTEM_KW = [
        'bilgi isteme yazısı',
        'bilgi talep formu',
        'bilgi talebine',
        'bilgi talebi',
        'karşıt inceleme bilgi talebi',
        'bilgi ve belge istemi',
        'bilgi istenmektedir',
    ]

    bilgi = {
        'karsit_firma_adi': '',
        'mukellef_adi':    '',   # tasdik verilen firma adı (2. ünvan = mükellef)
        'karsit_vkn': '',
        'mukellef_vkn': '',
        'kdv_soz_tarih': '',
        'kdv_soz_no': '',
        'tam_soz_tarih': '',
        'tam_soz_no': '',
        'tam_tasdik_var': False,
        'inceleme_donemi': '',
        'bilgi_istem': False,
    }
    _unvan_sayac = 0   # 1=YMM, 2=mükellef, 3=karşıt
    _vkn_sayac   = 0   # 1=YMM VKN (atla), 2=mükellef VKN, 3=karşıt VKN

    veri = doc_oku(word_yolu)
    if not veri:
        return bilgi

    tablolar = []
    if veri.get('html'):
        tablolar = html_tablolari_cek(veri['html'])
    elif veri.get('tablolar'):
        tablolar = veri['tablolar']

    # Tüm satırları düz liste olarak işle
    satirlar = [satir for tablo in tablolar for satir in tablo]

    # Dosya adına da bak
    dosya_adi_low = os.path.basename(word_yolu).lower()
    if any(kw in dosya_adi_low for kw in _BILGI_ISTEM_KW):
        bilgi['bilgi_istem'] = True

    for i, satir in enumerate(satirlar):
        if not satir:
            continue
        ilk = satir[0].strip() if satir else ''
        deger = satir[1].strip() if len(satir) > 1 else ''

        ilk_low = ilk.lower()

        # Konu satırında Bilgi İstem kontrolü
        if 'konu' in ilk_low and deger:
            deger_low = deger.lower()
            if any(kw in deger_low for kw in _BILGI_ISTEM_KW):
                bilgi['bilgi_istem'] = True

        # Sözleşme Dönemi- Konusu → Bilgi İsteme belgesi (İnceleme Dönemi- Konusu normal)
        if 'sözleşme dönemi' in ilk_low and 'konu' in ilk_low:
            bilgi['bilgi_istem'] = True

        # KDV Sözleşmesi Tarih – Sayısı
        if 'kdv sözleşme' in ilk_low or ('sözleşme' in ilk_low and 'tarih' in ilk_low and 'kdv' in ilk_low):
            # Değer: "02.02.2026 - 2026/006" veya "02.02.2026 – 2026/006"
            parca = re.split(r'\s*[-–]\s*', deger, maxsplit=1)
            if len(parca) >= 2:
                bilgi['kdv_soz_tarih'] = parca[0].strip()
                bilgi['kdv_soz_no']    = parca[1].strip()
            elif parca:
                bilgi['kdv_soz_tarih'] = parca[0].strip()

        # Tam Tasdik Sözleşmesi
        if 'tam tasdik' in ilk_low and 'sözleşme' in ilk_low:
            bilgi['tam_tasdik_var'] = True
            parca = re.split(r'\s*[-–]\s*', deger, maxsplit=1)
            if len(parca) >= 2:
                bilgi['tam_soz_tarih'] = parca[0].strip()
                bilgi['tam_soz_no']    = parca[1].strip()

        # İnceleme Dönemi - Konusu
        if 'inceleme dönemi' in ilk_low or 'dönem' in ilk_low and 'konu' in ilk_low:
            bilgi['inceleme_donemi'] = deger

        # VKN — "Vergi Dairesi / Vergi No" satırları
        # 1. oluşum = YMM VKN (atla), 2. = mükellef, 3. = karşıt
        if ('vergi dairesi' in ilk_low or 'vergi no' in ilk_low) and deger:
            vkn_match = re.search(r'\b(\d{10,11})\b', deger)
            if vkn_match:
                _vkn_sayac += 1
                vkn = vkn_match.group(1)
                if _vkn_sayac == 2 and not bilgi['mukellef_vkn']:
                    bilgi['mukellef_vkn'] = vkn
                elif _vkn_sayac >= 3 and not bilgi['karsit_vkn']:
                    bilgi['karsit_vkn'] = vkn

        # Firma adları — "Adı Soyadı / Ünvanı" satırları
        # 1. oluşum = YMM (atla), 2. = mükellef (tasdik verilen), 3. = karşıt
        if ('adı soyadı' in ilk_low or 'ünvan' in ilk_low) and deger:
            _unvan_sayac += 1
            if _unvan_sayac == 2 and not bilgi['mukellef_adi']:
                bilgi['mukellef_adi'] = deger
            elif _unvan_sayac >= 3 and not bilgi['karsit_firma_adi']:
                bilgi['karsit_firma_adi'] = deger

    return bilgi


def klasor_tara(word_klasor: str) -> list[dict]:
    """
    Word klasöründeki tüm .doc/.docx dosyalarını tarar.
    Döner: her dosya için {'dosya_yolu', 'firma_adi', ...tasdik_bilgisi_cek} listesi
    """
    import glob
    dosyalar = (
        glob.glob(os.path.join(word_klasor, '*.doc')) +
        glob.glob(os.path.join(word_klasor, '*.docx'))
    )
    dosyalar = [d for d in dosyalar if not os.path.basename(d).startswith('~$')]
    dosyalar = sorted(dosyalar)

    sonuclar = []
    for d in dosyalar:
        bilgi = tasdik_bilgisi_cek(d)
        # Firma adı yoksa dosya adından al
        if not bilgi['karsit_firma_adi']:
            bilgi['karsit_firma_adi'] = os.path.splitext(os.path.basename(d))[0]
            # Dosya adındaki tarih/saat damgasını temizle
            bilgi['karsit_firma_adi'] = re.sub(r'-\d{2}\.\d{2}\.\d{4}.*$', '', bilgi['karsit_firma_adi']).strip()
        bilgi['dosya_yolu'] = d
        sonuclar.append(bilgi)
    return sonuclar
